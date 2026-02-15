from django.shortcuts import render, redirect, get_object_or_404
from users.models import UserProfile, ActivityLog
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from .utils import generate_full_resume, enhance_with_ollama
from .models import resume
from io import BytesIO


@login_required
def generate_resume(request):
    """Generate a new AI-enhanced resume from the user's profile."""
    try:
        profile = request.user.userprofile
    except UserProfile.DoesNotExist:
        return redirect('edit_profile')

    template_choice = profile.resume_template_choice or 1

    try:
        enhanced = generate_full_resume(profile, request.user.email)
    except Exception:
        enhanced = {
            'summary': '',
            'skills': profile.skills,
            'education': profile.education,
            'experience': profile.experience,
            'projects': profile.projects,
            'achievements': profile.achievements,
        }

    # Save the generated resume
    new_resume = resume.objects.create(
        user=request.user,
        full_name=profile.full_name,
        phone_number=profile.phone_number,
        email=request.user.email,
        location=profile.location,
        summary=enhanced.get('summary', ''),
        skills=enhanced.get('skills', profile.skills),
        education=enhanced.get('education', profile.education),
        experience=enhanced.get('experience', profile.experience),
        projects=enhanced.get('projects', profile.projects),
        achievements=enhanced.get('achievements', profile.achievements),
        template_used=template_choice,
    )
    ActivityLog.objects.create(user=request.user, action='resume_generated')

    return redirect('view-resume', pk=new_resume.pk)


@login_required
def resume_details(request, pk):
    """View a single resume in the selected template."""
    my_resume = get_object_or_404(resume, pk=pk, user=request.user)

    template_choice = my_resume.template_used or 1
    template_name = f'resume_templates/template{template_choice}.html'

    context = {
        'resume': my_resume,
        'full_name': my_resume.full_name,
        'phone_number': my_resume.phone_number,
        'email': my_resume.email,
        'location': my_resume.location,
        'summary': my_resume.summary,
        'skills': my_resume.skills,
        'education': my_resume.education,
        'experience': my_resume.experience,
        'projects': my_resume.projects,
        'achievements': my_resume.achievements,
        'resume_id': my_resume.pk,
    }

    return render(request, template_name, context)


@login_required
def edit_resume(request, pk):
    """Edit a saved resume — all fields are editable inline."""
    my_resume = get_object_or_404(resume, pk=pk, user=request.user)

    if request.method == 'POST':
        my_resume.full_name = request.POST.get('full_name', my_resume.full_name)
        my_resume.phone_number = request.POST.get('phone_number', my_resume.phone_number)
        my_resume.email = request.POST.get('email', my_resume.email)
        my_resume.location = request.POST.get('location', my_resume.location)
        my_resume.summary = request.POST.get('summary', my_resume.summary)
        my_resume.skills = request.POST.get('skills', my_resume.skills)
        my_resume.education = request.POST.get('education', my_resume.education)
        my_resume.experience = request.POST.get('experience', my_resume.experience)
        my_resume.projects = request.POST.get('projects', my_resume.projects)
        my_resume.achievements = request.POST.get('achievements', my_resume.achievements)

        template_used = request.POST.get('template_used')
        if template_used:
            my_resume.template_used = int(template_used)

        my_resume.save()
        return redirect('view-resume', pk=my_resume.pk)

    context = {
        'resume': my_resume,
    }
    return render(request, 'resume_templates/edit_resume.html', context)


@login_required
def view_resume(request):
    """List all saved resumes."""
    allresume = resume.objects.filter(user=request.user)
    return render(request, 'resume_templates/view_resume.html', {'allresume': allresume})


@login_required
def delete_resume(request, pk):
    """Delete a saved resume."""
    resume.objects.filter(pk=pk, user=request.user).delete()
    return redirect('view')


@login_required
def download_resume_docx(request, pk=None):
    """Download a resume as DOCX. If pk provided, uses that resume; otherwise latest."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if pk:
        my_resume = get_object_or_404(resume, pk=pk, user=request.user)
    else:
        my_resume = resume.objects.filter(user=request.user).first()
        if not my_resume:
            return redirect('view')

    document = Document()

    # Style the document
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Name heading
    heading = document.add_heading(my_resume.full_name, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Contact info
    contact = document.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_run = contact.add_run(
        f"{my_resume.email}  •  {my_resume.phone_number}  •  {my_resume.location}"
    )
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Summary
    if my_resume.summary:
        document.add_heading('Professional Summary', level=1)
        document.add_paragraph(my_resume.summary)

    # Skills
    if my_resume.skills:
        document.add_heading('Skills', level=1)
        document.add_paragraph(my_resume.skills)

    # Education
    if my_resume.education:
        document.add_heading('Education', level=1)
        document.add_paragraph(my_resume.education)

    # Experience
    if my_resume.experience:
        document.add_heading('Experience', level=1)
        document.add_paragraph(my_resume.experience)

    # Projects
    if my_resume.projects:
        document.add_heading('Projects', level=1)
        document.add_paragraph(my_resume.projects)

    # Achievements
    if my_resume.achievements:
        document.add_heading('Achievements', level=1)
        document.add_paragraph(my_resume.achievements)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    safe_name = my_resume.full_name.replace(' ', '_')
    response['Content-Disposition'] = f'attachment; filename="{safe_name}_Resume.docx"'
    return response